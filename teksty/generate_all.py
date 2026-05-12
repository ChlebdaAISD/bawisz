#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generator dokumentów Word dla klienta Bawisz — przegląd tekstów wszystkich podstron.
Każdy .docx ma dwie kolumny: AKTUALNY TEKST | NOWY TEKST.
Klient wpisuje zmiany w kolumnie NOWY TEKST.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Kolory ─────────────────────────────────────────────────────────────
COLOR_HEADER_BG  = RGBColor(0x3A, 0x2E, 0x1F)
COLOR_SECTION_BG = RGBColor(0xA8, 0x80, 0x62)
COLOR_ROW_ALT    = RGBColor(0xF5, 0xF1, 0xEA)
COLOR_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TEXT       = RGBColor(0x3A, 0x2E, 0x1F)
COLOR_MUTED      = RGBColor(0x77, 0x62, 0x58)
COLOR_NOTE_BG    = RGBColor(0xFF, 0xF8, 0xE1)
COLOR_NOTE_TEXT  = RGBColor(0x7A, 0x5C, 0x00)
COLOR_CMS_BG     = RGBColor(0xEE, 0xF4, 0xFF)
COLOR_CMS_TEXT   = RGBColor(0x2E, 0x4E, 0x9A)

INSTRUKCJA = ("Instrukcja: W kolumnie NOWY TEKST wpisz zmienioną wersję "
              "(lub zostaw bez zmian). Nie modyfikuj kolumny AKTUALNY TEKST.")

OUTDIR = os.path.dirname(os.path.abspath(__file__))

# ── Helpers ─────────────────────────────────────────────────────────────
def _hex(color):
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"

def set_cell_bg(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color)); tcPr.append(shd)

def set_cell_border(cell, color='C8B89A', sz='4'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','bottom','left','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_cell_padding(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val)); m.set(qn('w:type'), 'dxa'); tcMar.append(m)
    tcPr.append(tcMar)

def add_label_para(cell, text):
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(text); r.font.size = Pt(7); r.font.bold = True
    r.font.color.rgb = COLOR_MUTED; r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)

def add_content_para(cell, text, bold=False, size=10):
    p = cell.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = COLOR_TEXT; r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

def style_header_row(table):
    row = table.rows[0]
    for i, label in enumerate(["AKTUALNY TEKST", "NOWY TEKST"]):
        cell = row.cells[i]
        set_cell_bg(cell, COLOR_HEADER_BG)
        set_cell_border(cell, color='3A2E1F', sz='6')
        set_cell_padding(cell, top=120, bottom=120)
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(label); r.font.size = Pt(9); r.font.bold = True
        r.font.color.rgb = COLOR_WHITE; r.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def new_doc(title):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(29.7); s.page_height = Cm(21.0)
    s.left_margin = s.right_margin = Cm(1.5)
    s.top_margin = s.bottom_margin = Cm(1.5)
    tp = doc.add_paragraph()
    r = tp.add_run(f"Bawisz — Przegląd tekstów: {title}")
    r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = COLOR_HEADER_BG; r.font.name = 'Calibri'
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(4)
    sp = doc.add_paragraph()
    sr = sp.add_run(INSTRUKCJA)
    sr.font.size = Pt(9); sr.font.italic = True
    sr.font.color.rgb = COLOR_MUTED; sr.font.name = 'Calibri'
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(10)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in t.columns:
        for cell in col.cells:
            cell.width = Cm(12.8)
    style_header_row(t)
    return doc, t

def sec(table, title):
    tr = table.add_row(); tr.cells[0].merge(tr.cells[1])
    cell = tr.cells[0]
    set_cell_bg(cell, COLOR_SECTION_BG)
    set_cell_border(cell, color='A88062', sz='8')
    set_cell_padding(cell, top=130, bottom=130, left=200, right=200)
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(title); r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = COLOR_WHITE; r.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

_row_counter = {'n': 0}
def row(table, label, text, bold=False, size=10):
    _row_counter['n'] += 1
    alt = (_row_counter['n'] % 2 == 0)
    tr = table.add_row()
    bg = COLOR_ROW_ALT if alt else COLOR_WHITE
    for ci in range(2):
        cell = tr.cells[ci]
        set_cell_bg(cell, bg)
        set_cell_border(cell, color='D8C8B0', sz='4')
        set_cell_padding(cell, top=100, bottom=100, left=160, right=160)
        add_label_para(cell, label)
        add_content_para(cell, text, bold=bold, size=size)

def reset_alt():
    _row_counter['n'] = 0

def note(table, text):
    tr = table.add_row(); tr.cells[0].merge(tr.cells[1])
    cell = tr.cells[0]
    set_cell_bg(cell, COLOR_NOTE_BG)
    set_cell_border(cell, color='E0C060', sz='4')
    set_cell_padding(cell, top=80, bottom=80, left=200, right=200)
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(f"⚠ {text}"); r.font.size = Pt(8); r.font.italic = True
    r.font.color.rgb = COLOR_NOTE_TEXT; r.font.name = 'Calibri'

def save(doc, path):
    doc.save(path)
    print(f"✅ {os.path.basename(path)}")

# ════════════════════════════════════════════════════════════════════════
# HOME — Strona główna
# ════════════════════════════════════════════════════════════════════════
def gen_home():
    reset_alt()
    doc, t = new_doc("Strona główna")

    # HERO
    sec(t, "1. HERO — banner powitalny")
    row(t, "EYEBROW (mały tekst nad nagłówkiem)", "[ Bawialnia Montessori · Nowy Targ ]")
    row(t, "NAGŁÓWEK H1 (linia 1)", "Bawialnia", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 2 — kursywa)", "w Nowym Targu,", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 3)", "do której dziecko", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 4 — kursywa)", "chce wracać.", bold=True, size=11)
    row(t, "LEAD (podtytuł, linia 1)", "Drewniana sala Montessori dla dzieci 0–10 lat przy ul. Krzywej 19B.")
    row(t, "LEAD (podtytuł, linia 2)", "Ty pijesz kawę i jesz domowe ciasto. Dziecko bawi się obok — przy naturalnych zabawkach.")
    row(t, "PRZYCISK 1", "Zarezerwuj urodziny")
    row(t, "PRZYCISK 2", "Co u nas znajdziesz")
    row(t, "META STAT 1 — liczba", "0–10")
    row(t, "META STAT 1 — etykieta", "lat")
    row(t, "META STAT 2 — liczba", "220 m²")
    row(t, "META STAT 2 — etykieta", "drewnianej sali")
    row(t, "META STAT 3 — liczba", "7 dni")
    row(t, "META STAT 3 — etykieta", "w tygodniu")

    # ABOUT
    sec(t, "2. O NAS — sekcja powitalna")
    row(t, "EYEBROW", "[ O nas ]")
    row(t, "NAGŁÓWEK H2", "Drewniana sala dla małych odkrywców.", bold=True, size=11)
    row(t, "LEAD", "Bawisz to 220 m² drewnianej przestrzeni Montessori w sercu Nowego Targu — sala zabaw dla dzieci 0–10 lat i kawiarnia dla rodziców pod jednym dachem. Dziecko bawi się samo wśród naturalnych zabawek, a ty masz chwilę przy dobrej kawie.")
    row(t, "WARTOŚĆ 1 — tytuł", "Naturalność", bold=True)
    row(t, "WARTOŚĆ 1 — opis", "Drewno, sklejka, tkaniny. Zabawki, którym kibicują rodzice — bez plastikowego hałasu i bez ekranów.")
    row(t, "WARTOŚĆ 2 — tytuł", "Spokój", bold=True)
    row(t, "WARTOŚĆ 2 — opis", "Dużo światła, mało bodźców. Pijesz kawę obok, dziecko bawi się w zasięgu wzroku.")
    row(t, "WARTOŚĆ 3 — tytuł", "Samodzielność", bold=True)
    row(t, "WARTOŚĆ 3 — opis", "Wszystko na wysokości dziecka. Wybiera samo, próbuje samo, wraca dumne z tego, co zrobiło.")
    row(t, "PRZYCISK CTA", "Czytaj więcej o nas")

    # OFERTA
    sec(t, "3. OFERTA — 4 karty (Sala / Kawiarnia / Urodziny / Warsztaty)")
    row(t, "EYEBROW", "[ Co u nas znajdziesz ]")
    row(t, "NAGŁÓWEK H2", "Cztery powody, żeby zostać na dłużej.", bold=True, size=11)
    row(t, "LEAD", "Bawialnia, kawiarnia, urodziny i warsztaty — wszystko w jednej drewnianej przestrzeni przy ul. Krzywej 19B w Nowym Targu. Wchodzisz na chwilę, zostajesz na trzy godziny.")
    row(t, "KARTA 01 — tytuł", "Sala zabaw", bold=True)
    row(t, "KARTA 01 — opis", "220 m² drewnianej przestrzeni Montessori — naturalne zabawki, strefy dla najmłodszych i starszych. Bawisz się ze swoim dzieckiem, nie zostawiasz go z animatorem.")
    row(t, "KARTA 01 — tagi", "0–10 lat · Montessori · Drewno")
    row(t, "KARTA 01 — CTA", "Czytaj więcej o Bawiszu")
    row(t, "KARTA 02 — tytuł", "Kawiarnia", bold=True)
    row(t, "KARTA 02 — opis", "Kawa specialty, domowe ciasta (też bez cukru i bezglutenowe), menu dla dzieci. Pijesz kawę obok sali — dziecko w zasięgu wzroku, bez wstawania.")
    row(t, "KARTA 02 — tagi", "Specialty · Domowe ciasta · Bez cukru")
    row(t, "KARTA 02 — CTA", "Zobacz menu kawiarni")
    row(t, "KARTA 03 — tytuł", "Urodziny", bold=True)
    row(t, "KARTA 03 — opis", "Pakiet MINI 45 zł/os. (2 h) albo STANDARD 74 zł/os. (2,5 h, sala tylko dla was). Dekoracje, poczęstunek i prezent dla solenizanta — my robimy resztę.")
    row(t, "KARTA 03 — tagi", "MINI · STANDARD · Sala na wyłączność · Prezent od nas")
    row(t, "KARTA 03 — CTA", "Zobacz pakiety urodzinowe")
    row(t, "KARTA 04 — tytuł", "Warsztaty", bold=True)
    row(t, "KARTA 04 — opis", "Plastyka, sensoplastyka, glina, joga dla dzieci. Małe grupy, prowadzące prowadzą zajęcia — ty siedzisz w kawiarni i masz chwilę.")
    row(t, "KARTA 04 — tagi", "Plastyka · Sensoplastyka · Małe grupy")
    row(t, "KARTA 04 — CTA", "Sprawdź warsztaty")

    # CENNIK
    sec(t, "4. CENNIK — pakiety czasowe i karnety")
    row(t, "EYEBROW", "[ Cennik ]")
    row(t, "NAGŁÓWEK H2", "Wstęp od 25 zł. Bez niespodzianek.", bold=True, size=11)
    row(t, "LEAD", "Cztery pakiety czasowe, jeden karnet miesięczny, karty podarunkowe od 29 zł. Opiekun zawsze gratis, rodzeństwo −25%.")
    row(t, "PAKIET 1 — czas", "1h")
    row(t, "PAKIET 1 — cena", "25 zł")
    row(t, "PAKIET 2 — czas", "1,5 h")
    row(t, "PAKIET 2 — cena", "29 zł")
    row(t, "PAKIET 3 — czas", "2 h")
    row(t, "PAKIET 3 — cena", "33 zł")
    row(t, "PAKIET 4 — czas", "NO LIMIT")
    row(t, "PAKIET 4 — cena", "45 zł")
    row(t, "ZNIŻKA — opis", "Zniżki dla rodzeństwa −25% (drugie i kolejne dziecko)")
    row(t, "KARNET — tytuł", "Karnet miesięczny", bold=True)
    row(t, "KARNET — podtytuł", "no limit · cały miesiąc kalendarzowy")
    row(t, "KARNET — cena", "215 zł")
    row(t, "KARTA PODAR. — tytuł", "Karta podarunkowa", bold=True)
    row(t, "KARTA PODAR. 1", "1 wejście × 1,5 h — 29 zł")
    row(t, "KARTA PODAR. 2", "3 wejścia × 1,5 h — 75 zł")
    row(t, "KARTA PODAR. 3", "5 wejść × 1,5 h — 135 zł")
    row(t, "PRZYCISK CTA", "Zadzwoń · 693 766 049")

    # URODZINY
    sec(t, "5. URODZINY — pakiety MINI i STANDARD")
    row(t, "EYEBROW", "[ Urodziny ]")
    row(t, "NAGŁÓWEK H2", "Oferta urodzinowa", bold=True, size=11)
    row(t, "LEAD", "Dwa pakiety, jedno wspomnienie na lata. Cały dekor, poczęstunek i opieka po naszej stronie — wy bawicie się z dzieckiem.")
    row(t, "PAKIET MINI — nazwa", "MINI", bold=True)
    row(t, "PAKIET MINI — czas", "2 h")
    row(t, "PAKIET MINI — cena", "45 zł / os.")
    row(t, "PAKIET MINI — poczęstunek 1", "woda no limit")
    row(t, "PAKIET MINI — poczęstunek 2", "sok tłoczony 200 ml/os.")
    row(t, "PAKIET MINI — poczęstunek 3", "paluszki")
    row(t, "PAKIET MINI — poczęstunek 4", "OTO chrupki")
    row(t, "PAKIET MINI — poczęstunek 5", "galaretki")
    row(t, "PAKIET MINI — dekoracje 1", "kolorowa zastawa lub w naturalnym stylu")
    row(t, "PAKIET MINI — dekoracje 2", "balon cyfra")
    row(t, "PAKIET STANDARD — badge", "Najczęściej wybierany")
    row(t, "PAKIET STANDARD — nazwa", "STANDARD", bold=True)
    row(t, "PAKIET STANDARD — czas", "2,5 h")
    row(t, "PAKIET STANDARD — cena", "74 zł / os.")
    row(t, "PAKIET STANDARD — perk 1", "sala na wyłączność")
    row(t, "PAKIET STANDARD — perk 2", "minimum 10 dzieci")
    row(t, "PAKIET STANDARD — perk 3", "prezent dla solenizanta")
    row(t, "PAKIET STANDARD — poczęstunek 1", "woda no limit")
    row(t, "PAKIET STANDARD — poczęstunek 2", "sok tłoczony no limit")
    row(t, "PAKIET STANDARD — poczęstunek 3", "owoce · paluszki")
    row(t, "PAKIET STANDARD — poczęstunek 4", "OTO chrupki · gofry")
    row(t, "PAKIET STANDARD — poczęstunek 5", "cake pops lub babeczki · galaretki")
    row(t, "PAKIET STANDARD — dekoracje 1", "zastawa stołowa w wybranym motywie")
    row(t, "PAKIET STANDARD — dekoracje 2", "ścianka dekoracyjna")
    row(t, "PAKIET STANDARD — dekoracje 3", "balon cyfra")
    row(t, "PAKIET STANDARD — dekoracje 4", "girlanda balonowa na sali zabaw")
    row(t, "PAKIET STANDARD — dekoracje 5", "cyfrowe zaproszenia w danej tematyce")

    # MENU
    sec(t, "6. KAWIARNIA / MENU — kawa, ciasta, dla dzieci")
    row(t, "EYEBROW", "[ Kawiarnia ]")
    row(t, "NAGŁÓWEK H2", "Kawa, dla której tu wracasz.", bold=True, size=11)
    row(t, "LEAD", "Specialty kawa, domowe ciasta (też bez cukru i bezglutenowe), menu dla dzieci bez cukru. Mała kawiarnia obok sali zabaw — pijesz, jesz, dziecko się bawi w zasięgu wzroku.")
    row(t, "ZAKŁADKA 1", "Kawa & napoje")
    row(t, "ZAKŁADKA 2", "Ciasta")
    row(t, "ZAKŁADKA 3", "Dla dzieci")
    row(t, "KAWA — pozycja 1", "Espresso — mocne i krótkie, jak trzeba")
    row(t, "KAWA — pozycja 2", "Flat white — jedwabista mleczna pianka")
    row(t, "KAWA — pozycja 3", "Latte — z syropem waniliowym, orzechowym lub karmelowym")
    row(t, "KAWA — pozycja 4", "Cappuccino — klasyka, której nie trzeba przedstawiać")
    row(t, "KAWA — pozycja 5", "Matcha latte — mleko zwykłe lub roślinne")
    row(t, "KAWA — pozycja 6", "Czekolada na gorąco — rozgrzewa lepiej niż piec kaflowy")
    row(t, "CIASTA — pozycja 1", "Sezonowe ciasta domowe — pieczone u nas, zmieniają się co tydzień")
    row(t, "CIASTA — pozycja 2", "Brownie — ciemna czekolada, na zimno albo lekko podgrzane")
    row(t, "CIASTA — pozycja 3", "Sernik — klasyczny lub z sezonowym dodatkiem")
    row(t, "CIASTA — pozycja 4", "Szarlotka — jabłka, cynamon, opcjonalnie z lodami")
    row(t, "CIASTA — pozycja 5", "Wersje bez cukru i bezglutenowe — pytaj na miejscu")
    row(t, "DLA DZIECI — pozycja 1", "Owoce sezonowe — prosto z talerza, krojone na kawałki")
    row(t, "DLA DZIECI — pozycja 2", "Mleko / kakao — ciepłe albo zimne")
    row(t, "DLA DZIECI — pozycja 3", "Soczki naturalne — bez dodatku cukru")
    row(t, "DLA DZIECI — pozycja 4", "Babeczki bananowe — bez cukru, zawsze świeże")
    row(t, "DLA DZIECI — pozycja 5", "Kanapka z dżemem — prosto i pewnie")
    row(t, "STOPKA MENU", "* Pełne, aktualne menu z cenami znajdziesz na miejscu. Oferta zmienia się sezonowo.")
    row(t, "PODPIS ZDJĘCIA — tytuł", "Kawa i ciasto")
    row(t, "PODPIS ZDJĘCIA — podtytuł", "Tymczasem dzieci bawią się tuż obok.")
    row(t, "PRZYCISK CTA", "Sprawdź kawiarnię")

    # TESTIMONIALS
    sec(t, "7. OPINIE RODZICÓW")
    row(t, "EYEBROW", "[ Co mówią rodzice ]")
    row(t, "OCENA — liczba", "4.9", bold=True)
    row(t, "OCENA — podpis", "ocena w Google")
    row(t, "OPINIA 1 — autor", "Katarzyna Domańska", bold=True)
    row(t, "OPINIA 1 — źródło", "opinia z Google")
    row(t, "OPINIA 1 — treść", "BAWISZ to najlepsza bawialnia, w jakiej byłam z moją 3-latką. Spędziłam tu cały dzień: ja siedząc w kawiarni z pysznym ciastem i kawą, mając całą bawialnię w zasięgu wzroku, a córka bawiąc się w pięknym, dopracowanym w każdym detalu placu zabaw.")
    row(t, "OPINIA 2 — autor", "Anita Słodyczka", bold=True)
    row(t, "OPINIA 2 — źródło", "opinia z Google")
    row(t, "OPINIA 2 — treść", "Świetne miejsce! Wszystko drewniane łącznie z zabawkami — to ogromny plus. Córka bawiła się tak dobrze, że nie chciała wyjść. Czystość, przepiękne wnętrza, pyszna kawa, herbata i ciasto. Polecam gorąco!")
    row(t, "OPINIA 3 — autor", "Karolina Biela", bold=True)
    row(t, "OPINIA 3 — źródło", "opinia z Google")
    row(t, "OPINIA 3 — treść", "Gorąco polecam dla dziecka. Moja 9-miesięczna córka była zachwycona zabawkami, bezpiecznymi i różnorodnymi. Dziecko może się pobawić, a rodzic odpocząć przy pysznej kawie, nie spuszczając go z oczu. Obsługa miła i pomocna.")
    row(t, "OPINIA 4 — autor", "Adrian T.", bold=True)
    row(t, "OPINIA 4 — źródło", "opinia z Google")
    row(t, "OPINIA 4 — treść", "Super miejsce dla dzieci! Sala zabaw to strzał w 10 — świetnie zorganizowana i naprawdę fajna przestrzeń dla najmłodszych. Dodatkowo monoporcje są przepyszne. Bardzo polecam.")
    row(t, "OPINIA 5 — autor", "Beata Waras", bold=True)
    row(t, "OPINIA 5 — źródło", "opinia z Google")
    row(t, "OPINIA 5 — treść", "Super miejsce dla maluszka — dużo pięknych i interesujących zabawek, pyszna kawa. Dodatkowa gwiazdka za wyposażenie toalety dla bobasów: pieluszki, mokre chusteczki, nocnik, nakładka. Polecam!")

    # GODZINY
    sec(t, "8. GODZINY OTWARCIA + ADRES")
    row(t, "EYEBROW", "[ Godziny ]")
    row(t, "NAGŁÓWEK H2", "Codziennie otwarte.", bold=True, size=11)
    row(t, "GODZINY — Pn-Pt", "Poniedziałek–Piątek: 10:00 — 19:00")
    row(t, "GODZINY — Sob-Nd", "Sobota–Niedziela: 10:00 — 20:00 (dłużej)")
    row(t, "KARTA ADRES — eyebrow", "[ Jak do nas trafić? ]")
    row(t, "KARTA ADRES — nagłówek", "ul. Krzywa 19B, Nowy Targ", bold=True)
    row(t, "ADRES — kod pocztowy + miasto", "34-400 Nowy Targ")
    row(t, "ADRES — opis", "Centrum, parking obok")
    row(t, "TELEFON", "+48 693 766 049")
    row(t, "TELEFON — podpis", "Zadzwoń lub napisz SMS")
    row(t, "GODZINY — krótko (tytuł)", "Codziennie od 10:00", bold=True)
    row(t, "GODZINY — krótko (opis)", "Pn–Pt do 19, Sb–Nd do 20")
    row(t, "PRZYCISK CTA", "Nawiguj do nas")

    # CONTACT
    sec(t, "9. KONTAKT — telefon i social media")
    row(t, "EYEBROW", "[ Kontakt ]")
    row(t, "NAGŁÓWEK H2 (linia 1)", "Pytanie?", bold=True, size=11)
    row(t, "NAGŁÓWEK H2 (linia 2)", "Zadzwoń", bold=True, size=11)
    row(t, "NAGŁÓWEK H2 (linia 3 — kursywa)", "albo napisz.", bold=True, size=11)
    row(t, "LEAD", "Rezerwacja urodzin, pytania o warsztaty, terminy wynajmu sali — najszybciej złapiesz nas telefonicznie. Odpisujemy też na Messengerze i Instagramie.")
    row(t, "PRZYCISK TEL", "+48 693 766 049")
    row(t, "PRZYCISK CTA 2", "Pełny kontakt")
    row(t, "PODTYTUŁ SOCIAL", "Albo śledź nas tutaj.")
    row(t, "INSTAGRAM — nazwa", "@bawisz_bawialnia")
    row(t, "INSTAGRAM — podpis", "Codzienne kadry")
    row(t, "FACEBOOK — nazwa", "Bawisz · bawialnia")
    row(t, "FACEBOOK — podpis", "Wydarzenia i info")
    row(t, "TIKTOK — nazwa", "@bawisz.bawialnia")
    row(t, "TIKTOK — podpis", "Kulisy i zabawy")

    # FAQ
    sec(t, "10. FAQ — najczęstsze pytania")
    home_faq = [
        ('Ile kosztuje wstęp do bawialni w Nowym Targu?',
         '1 godzina to 25 zł, 1,5 h — 29 zł, 2 h — 33 zł, a cały dzień (NO LIMIT) — 45 zł. Karnet miesięczny NO LIMIT (cały miesiąc kalendarzowy) kosztuje 215 zł. Opiekun zawsze gratis, drugie i każde kolejne dziecko z rodzeństwa ma 25% zniżki.'),
        ('Dla jakiego wieku jest bawialnia Bawisz?',
         'Sala jest przemyślana dla dzieci od 0 do 10 lat. Drewniana przestrzeń w duchu Montessori dzieli się na strefy dopasowane do wieku — niemowlę, dwulatek i ośmiolatek znajdą u nas coś dla siebie.'),
        ('Czym Bawisz różni się od zwykłego placu zabaw?',
         'Trzy konkrety. Po pierwsze — drewniany wystrój sali i naturalne zabawki, bez plastikowego hałasu. Po drugie — kawiarnia obok, więc pijesz prawdziwą kawę i jesz domowe ciasto, kiedy dziecko się bawi. Po trzecie — w codziennym wstępie nie ma animatorów, bawisz się ze swoim dzieckiem w przestrzeni, w której wszystko jest pod ręką.'),
        ('Czy trzeba rezerwować wejście do bawialni?',
         'Na codzienny wstęp nie. Wpadasz, kiedy chcesz, w godzinach otwarcia (10:00 — 19:00 w tygodniu, 10:00 — 20:00 w weekendy). Rezerwujemy tylko urodziny, warsztaty i wyjścia grupowe dla przedszkoli — wtedy zadzwoń pod +48 693 766 049 albo napisz na Instagramie.'),
        ('Czy jest parking przy bawialni w Nowym Targu?',
         'Tak — parkujesz tuż obok lokalu, w centrum miasta. Adres: ul. Krzywa 19B, 34-400 Nowy Targ. To dwie minuty pieszo od rynku.'),
        ('Co to jest bawialnia Montessori i dlaczego to ważne?',
         'W Montessori dziecko wybiera samo, próbuje samo, jest dumne z efektu. Wszystko jest na jego wysokości — drewniane konstrukcje, sklejka, tkaniny, zabawki bez baterii i ekranów. Mniej hałasu, więcej skupienia. Dzieci wciągają się tak, że trudno je wyciągnąć po dwóch godzinach.'),
        ('Czy mogę zostawić dziecko pod opieką personelu?',
         'Nie — i to jest świadome. Bawisz to przestrzeń, w której bawisz się ze swoim dzieckiem, a nie miejsce, gdzie zostawiasz dziecko z animatorem. Sala jest dograna pod dziecko, więc nie musisz pilnować na każdym kroku — siedzisz z kawą obok, dziecko ma wszystko w zasięgu ręki. Wyjątek to urodziny i warsztaty — tam prowadzący prowadzą program.'),
        ('Czy są zniżki dla rodzeństwa?',
         'Tak — drugie i każde kolejne dziecko z rodzeństwa ma 25% zniżki na każdy pakiet czasowy. Opiekun (rodzic, dziadek, babcia) zawsze wchodzi gratis. Mamy też karty podarunkowe od 29 zł — popularny prezent na chrzciny i pierwsze urodziny.'),
    ]
    row(t, "EYEBROW", "[ FAQ ]")
    row(t, "NAGŁÓWEK H2", "Najczęstsze pytania.", bold=True, size=11)
    for i, (q, a) in enumerate(home_faq, 1):
        row(t, f"PYTANIE {i}", q, bold=True)
        row(t, f"ODPOWIEDŹ {i}", a)

    # META
    sec(t, "11. META — SEO (Title + Description)")
    note(t, "Title i Description są kluczowe dla Google. Title 50-60 znaków, Description 150-160 znaków.")
    row(t, "META TITLE", "Bawialnia Nowy Targ — Bawisz | drewniana sala Montessori")
    row(t, "META DESCRIPTION", "Bawialnia Montessori w Nowym Targu dla dzieci 0-10 lat. Drewniana sala 220 m², kawiarnia obok. Wstęp od 25 zł, urodziny od 45 zł/os. Ocena 4.9/5 w Google.")

    save(doc, os.path.join(OUTDIR, "home.docx"))


# ════════════════════════════════════════════════════════════════════════
# URODZINY
# ════════════════════════════════════════════════════════════════════════
def gen_urodziny():
    reset_alt()
    doc, t = new_doc("Urodziny")

    sec(t, "1. HERO — banner powitalny")
    row(t, "BREADCRUMB", "Strona główna › Urodziny")
    row(t, "NAGŁÓWEK H1 (linia 1)", "Urodziny dla dziecka.", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 2 — kursywa)", "Nowy Targ — drewniana", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 3)", "bawialnia Montessori.", bold=True, size=11)
    row(t, "PRZYCISK 1", "Zarezerwuj urodziny")
    row(t, "PRZYCISK 2", "Zadzwoń · 693 766 049")

    sec(t, "2. INTRO — direct answer (pierwsze 100 słów)")
    row(t, "AKAPIT INTRO", "Urodziny w Bawiszu to drewniana sala w duchu Montessori przy ul. Krzywej 19B w Nowym Targu, dwa pakiety (MINI 45 zł/os. za 2 godziny, STANDARD 74 zł/os. za 2,5 godziny z salą tylko dla was i minimum 10 dzieci), pełne dekoracje, poczęstunek i kawiarnia obok. Tort przynosisz ty, resztę robimy my. Sala jest dla dzieci od 0 do 10 lat — nikt nie zabiera ci dziecka, bawicie się razem w przestrzeni, w której wszystko jest dograne pod dziecko.")

    sec(t, "3. PAKIETY (MINI i STANDARD) — patrz home.docx sekcja 5")
    note(t, "Pakiety MINI i STANDARD są opisane w pliku home.docx (sekcja 5 — URODZINY). Klient może edytować je raz w pliku home.docx — będą zaktualizowane wszędzie.")

    sec(t, "4. GALERIA — podpisy zdjęć")
    row(t, "EYEBROW", "[ Galeria ]")
    row(t, "NAGŁÓWEK H2", "Tak wyglądają urodziny w Bawiszu.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Drewniana sala, naturalne zabawki, dekoracje motywowe. Kliknij zdjęcie, żeby powiększyć.")
    row(t, "ALT zdjęcia 1", "Sala urodzinowa w Bawiszu")
    row(t, "ALT zdjęcia 2", "Strefa zabaw na urodzinach")
    row(t, "ALT zdjęcia 3", "Dzieci na urodzinach")
    row(t, "ALT zdjęcia 4", "Stół urodzinowy")
    row(t, "ALT zdjęcia 5", "Dekoracje urodzinowe")

    sec(t, "5. PROCES — 4 kroki")
    row(t, "EYEBROW", "[ Jak to wygląda ]")
    row(t, "NAGŁÓWEK H2", "Cztery kroki. Bez niespodzianek.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Najczęstsze pytanie: „co mam zrobić, a co robicie wy?\". Odpowiedź: ty przynosisz tort i gości, resztę robimy my.")
    row(t, "KROK 1 — tytuł", "Telefon i ustalenie terminu", bold=True)
    row(t, "KROK 1 — opis", "Dzwonisz na +48 693 766 049 albo piszesz na Instagramie. Wspólnie wybieramy dzień, godzinę, pakiet (MINI lub STANDARD), liczbę dzieci i motyw dekoracji (przy pakiecie STANDARD).")
    row(t, "KROK 2 — tytuł", "Przygotowanie sali", bold=True)
    row(t, "KROK 2 — opis", "Zanim przyjdziecie, ustawiamy stoły, dekoracje, balon-cyfrę i poczęstunek. W pakiecie STANDARD dostajesz cyfrowe zaproszenie, żeby rozesłać je gościom.")
    row(t, "KROK 3 — tytuł", "Bawicie się razem", bold=True)
    row(t, "KROK 3 — opis", "Dzieci bawią się z tobą przy naturalnych zabawkach Montessori (drewno, sensoplastyka, zabawy w role). Bez animatora i wyreżyserowanych zabaw. Sala jest przemyślana — bawicie się spokojnie, bo wszystko jest bezpieczne i pod ręką.")
    row(t, "KROK 4 — tytuł", "Tort i pamiątkowe zdjęcia", bold=True)
    row(t, "KROK 4 — opis", "Tort solenizanta przynosicie wy — my serwujemy i pomagamy przy świeczkach. W pakiecie STANDARD solenizant dostaje od nas prezent. Zdjęcia, jedzenie, świętowanie.")

    sec(t, "6. MID-CTA — środkowy banner")
    row(t, "NAGŁÓWEK H2", "Termin urodzin?", bold=True, size=11)
    row(t, "OPIS", "Pakiet STANDARD rezerwuje się 3-6 tygodni wcześniej. Najszybciej przez telefon — od razu sprawdzamy wolne daty i ustalamy motyw dekoracji.")
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "Napisz na Instagramie")

    sec(t, "7. FAQ — najczęstsze pytania")
    urodziny_faq = [
        ('Ile kosztują urodziny dla dziecka w Nowym Targu w Bawiszu?',
         'Pakiet MINI to 45 zł od osoby (2 godziny, poczęstunek, dekoracje, balon-cyfra). Pakiet STANDARD to 74 zł od osoby (2,5 godziny, sala tylko dla was, prezent dla solenizanta, pełna dekoracja tematyczna, cyfrowe zaproszenia). Minimum dla pakietu STANDARD to 10 dzieci.'),
        ('Jak długo trwają urodziny i co jest w cenie?',
         'MINI trwa 2 godziny, STANDARD 2,5 godziny. W obu pakietach dostajesz pełny poczęstunek (woda, sok tłoczony, paluszki, chrupki OTO, galaretki — w STANDARD dodatkowo gofry, owoce i cake pops lub babeczki) oraz dekorację stołu. STANDARD obejmuje też salę tylko dla was, ściankę dekoracyjną, girlandę balonową i prezent dla solenizanta.'),
        ('Dla jakiego wieku dziecka są te urodziny?',
         'Sala jest dla dzieci od 0 do 10 lat. Drewniane wnętrze w duchu Montessori dzieli się na strefy dopasowane do wieku: młodsze dzieci (2-4 lata) bawią się w innej części niż starsze (5-10 lat).'),
        ('Czy mogę przynieść własny tort?',
         'Tak — tort solenizanta przynosisz ty. Zapewniamy talerzyki, świeczki, stolik na tort i pomoc przy krojeniu. Reszta poczęstunku (paluszki, owoce, gofry, chrupki, galaretki) jest po naszej stronie.'),
        ('Czy rodzice biorą udział w urodzinach?',
         'Tak — to dla nas ważne. W Bawiszu dzieci bawią się razem z rodzicami przy naturalnych zabawkach Montessori. Nie zostawiasz dziecka pod opieką obsługi. My zajmujemy się dekoracjami, poczęstunkiem i całą logistyką, ty bawisz się ze swoim dzieckiem.'),
        ('Z jakim wyprzedzeniem trzeba rezerwować termin?',
         'Pakiet STANDARD (sala tylko dla was) — minimum 3-4 tygodnie wcześniej, w sezonie (kwiecień-czerwiec, listopad-grudzień) nawet 6 tygodni. Pakiet MINI udaje się czasem zarezerwować w tym samym tygodniu. Najszybciej idzie przez telefon: +48 693 766 049.'),
    ]
    row(t, "EYEBROW", "[ FAQ ]")
    row(t, "NAGŁÓWEK H2", "Najczęstsze pytania.", bold=True, size=11)
    for i, (q, a) in enumerate(urodziny_faq, 1):
        row(t, f"PYTANIE {i}", q, bold=True)
        row(t, f"ODPOWIEDŹ {i}", a)

    sec(t, "8. FINAL CTA — końcowy banner")
    row(t, "NAGŁÓWEK H2", "Rezerwacja urodzin", bold=True, size=11)
    row(t, "OPIS", "Pakiet MINI od 45 zł/os. albo STANDARD 74 zł/os. z salą tylko dla was. Decyzję podejmujesz, kiedy ustalimy datę i motyw dekoracji.")
    row(t, "PRZYCISK 1", "Zarezerwuj urodziny")
    row(t, "PRZYCISK 2", "Zadzwoń · 693 766 049")

    sec(t, "9. META — SEO")
    row(t, "META TITLE", "Urodziny dla dziecka Nowy Targ — Bawisz | Pakiety MINI 45 zł, STANDARD 74 zł")
    row(t, "META DESCRIPTION", "Urodziny dla dziecka w Nowym Targu w drewnianej bawialni Montessori. Pakiet MINI od 45 zł/os., STANDARD od 74 zł/os. z salą tylko dla was. Ocena 4.9/5 w Google.")

    save(doc, os.path.join(OUTDIR, "urodziny.docx"))


# ════════════════════════════════════════════════════════════════════════
# KAWIARNIA
# ════════════════════════════════════════════════════════════════════════
def gen_kawiarnia():
    reset_alt()
    doc, t = new_doc("Kawiarnia")

    sec(t, "1. HERO — banner powitalny")
    row(t, "BREADCRUMB", "Strona główna › Kawiarnia")
    row(t, "NAGŁÓWEK H1 (linia 1)", "Kawiarnia w Nowym Targu.", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 2 — kursywa)", "Z drewnianą bawialnią", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 3)", "Montessori obok.", bold=True, size=11)
    row(t, "PRZYCISK 1", "Zobacz menu")
    row(t, "PRZYCISK 2", "Zadzwoń · 693 766 049")

    sec(t, "2. INTRO — direct answer (pierwsze 100 słów)")
    row(t, "AKAPIT INTRO", "Kawiarnia w Nowym Targu, w której kawa zostaje gorąca, a dziecko nie ciągnie cię za rękę, żeby już wracać do domu. Przy ul. Krzywej 19B parzymy specialty espresso, flat white, latte i matcha latte, pieczemy ciasta domowe (też bez cukru i bezglutenowe), a obok mamy drewnianą salę Montessori dla dzieci od 0 do 10 lat. Pijesz kawę, dziecko bawi się tuż obok — przy naturalnych zabawkach, w spokojnej przestrzeni, bez plastiku. Wstęp do bawialni od 25 zł za godzinę, sama kawiarnia bez biletu. Ocena 4.9/5 w Google.")

    sec(t, "3. MENU (Kawa / Ciasta / Dla dzieci) — patrz home.docx sekcja 6")
    note(t, "Menu kawiarni jest opisane w pliku home.docx (sekcja 6). Edytuj raz tam — zmiana będzie wszędzie.")

    sec(t, "4. GALERIA — podpisy zdjęć")
    row(t, "EYEBROW", "[ Galeria ]")
    row(t, "NAGŁÓWEK H2", "Tak wygląda kawiarnia w Bawiszu.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Specialty coffee, domowe ciasta, dziecko bawi się obok. Kliknij zdjęcie, żeby powiększyć.")
    row(t, "ALT zdjęcia 1", "Kawiarnia w Bawiszu")
    row(t, "ALT zdjęcia 2", "Domowe ciasta")
    row(t, "ALT zdjęcia 3", "Specialty coffee")
    row(t, "ALT zdjęcia 4", "Wnętrze kawiarni")
    row(t, "ALT zdjęcia 5", "Menu dla dzieci")

    sec(t, "5. PROCES — 4 kroki (jak to działa)")
    row(t, "EYEBROW", "[ Jak to wygląda ]")
    row(t, "NAGŁÓWEK H2", "Cztery kroki. Od kawy po dziecko, które nie chce wychodzić.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Najczęstsze pytanie: „muszę kupować bilet, żeby napić się kawy?\". Odpowiedź: nie. Bilet jest tylko jeśli dziecko ma się bawić w sali.")
    row(t, "KROK 1 — tytuł", "Wpadasz na kawę", bold=True)
    row(t, "KROK 1 — opis", "Bez rezerwacji. Wchodzisz, wybierasz kawę i ciasto z lady (sezonowe domowe, brownie, sernik, szarlotka), rozsiadasz się przy stoliku. Do kawiarni nie potrzebujesz biletu — kupujesz go tylko jeśli dziecko ma się bawić w sali.")
    row(t, "KROK 2 — tytuł", "Decydujesz, czy bawialnia", bold=True)
    row(t, "KROK 2 — opis", "Jeśli dziecko jest z tobą i ma się pobawić — kupujesz wstęp do drewnianej sali Montessori (25 zł / 1 h, 29 zł / 1,5 h, 33 zł / 2 h, 45 zł no limit). Opiekun zawsze gratis. Drugie dziecko w rodzinie — 25% taniej.")
    row(t, "KROK 3 — tytuł", "Pijesz kawę, dziecko się bawi", bold=True)
    row(t, "KROK 3 — opis", "Siedzisz obok sali zabaw, masz dziecko w polu widzenia, kawa zostaje gorąca. Dziecko bawi się przy naturalnych zabawkach z drewna — sensoplastyka, zabawy w role, kącik czytelniczy.")
    row(t, "KROK 4 — tytuł", "Wracasz, kiedy chcesz", bold=True)
    row(t, "KROK 4 — opis", "Karta podarunkowa (3 wejścia × 1,5 h za 75 zł, 5 wejść × 1,5 h za 135 zł) albo karnet miesięczny no limit za 215 zł — jeśli przewidujesz, że to nie ostatni raz. Najprościej dogadać to przy ladzie.")

    sec(t, "6. MID-CTA — środkowy banner")
    row(t, "NAGŁÓWEK H2", "Zamówienie na imprezę?", bold=True, size=11)
    row(t, "OPIS", "Cały sernik, blacha brownie, ciasta na chrzciny albo urodziny w domu — ustalamy z minimum 24-48 h wyprzedzeniem. Najszybciej przez telefon.")
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "Napisz na Instagramie")

    sec(t, "7. FAQ — najczęstsze pytania")
    kawiarnia_faq = [
        ('Czy mogę przyjść z dzieckiem na kawę i mieć gdzie je posadzić?',
         'Tak — to po to powstał Bawisz. Pijesz kawę w kawiarni przy ul. Krzywej 19B, a dziecko obok bawi się w drewnianej sali Montessori. Wstęp do bawialni: od 25 zł za godzinę, opiekun zawsze gratis. Nie musisz żonglować dzieckiem na kolanach — siedzisz spokojnie, a dziecko ma wszystko, czego potrzebuje, pod ręką.'),
        ('Jakie ciasta pieczecie na miejscu?',
         'Sezonowe ciasta domowe (zmieniają się co tydzień), brownie, sernik i szarlotkę. Mamy też wersje bez cukru i bezglutenowe — zapytaj na miejscu, bo dostępność zależy od dnia. Wszystko pieczemy u nas, nie zamawiamy z hurtowni.'),
        ('Jaką kawę serwujecie?',
         'Espresso, flat white, latte (z syropem waniliowym, orzechowym lub karmelowym), cappuccino, matcha latte (mleko zwykłe lub roślinne), gorącą czekoladę. Pełne menu z cenami jest na miejscu — oferta zmienia się sezonowo.'),
        ('Czy macie coś dla dziecka do jedzenia i picia?',
         'Mamy: krojone owoce sezonowe, ciepłe lub zimne mleko i kakao, soczki naturalne bez dodatku cukru, babeczki bananowe bez cukru i kanapki z dżemem. Dla dziecka, które ma alergię albo dietę bez cukru, znajdziesz opcję, której nie ma w sieciówce.'),
        ('Czy muszę kupować bilet do bawialni, żeby zjeść u was ciasto?',
         'Nie — kawiarnia jest dostępna bez biletu do bawialni. Wpadasz na kawę i ciasto, siadasz, jesz. Bilet kupujesz tylko jeśli dziecko ma się bawić w sali Montessori. Niejedna mama przychodzi rano sama na kawę i wraca po południu już z dzieckiem.'),
        ('Gdzie was szukać i czy jest parking?',
         'ul. Krzywa 19B, 34-400 Nowy Targ — w centrum miasta. Parking jest tuż obok lokalu. Otwarte codziennie: pon.-pt. 10:00-19:00, sob.-niedz. 10:00-20:00. Dojazd z Krakowa, Zakopanego, Rabki czy Bukowiny Tatrzańskiej bez krążenia po mieście.'),
        ('Czy mogę zamówić ciasto na wynos albo na zamówienie?',
         'Ciasta z lady bierzesz na wynos bez problemu — pakujemy w pudełko. Większe zamówienia (cały sernik, blacha brownie, ciasta na imprezę domową) ustalamy telefonicznie z minimum 24-48 h wyprzedzeniem. Zadzwoń na +48 693 766 049 i powiedz, na kiedy i ile potrzebujesz.'),
    ]
    row(t, "EYEBROW", "[ FAQ ]")
    row(t, "NAGŁÓWEK H2", "Najczęstsze pytania.", bold=True, size=11)
    for i, (q, a) in enumerate(kawiarnia_faq, 1):
        row(t, f"PYTANIE {i}", q, bold=True)
        row(t, f"ODPOWIEDŹ {i}", a)

    sec(t, "8. FINAL CTA — końcowy banner")
    row(t, "NAGŁÓWEK H2", "Wpadasz na kawę?", bold=True, size=11)
    row(t, "OPIS", "ul. Krzywa 19B, Nowy Targ. Otwarte codziennie: pon.-pt. 10:00-19:00, sob.-niedz. 10:00-20:00. Bez rezerwacji — przyjdź, kiedy chcesz. Jeśli planujesz urodziny dziecka, zarezerwuj salę z wyprzedzeniem.")
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "Zarezerwuj urodziny")

    sec(t, "9. META — SEO")
    row(t, "META TITLE", "Kawiarnia Nowy Targ — Bawisz | specialty + domowe ciasta")
    row(t, "META DESCRIPTION", "Kawiarnia w Nowym Targu z bawialnią Montessori obok. Specialty kawa, domowe ciasta (też bez cukru i bezglutenowe). ul. Krzywa 19B. Ocena 4.9/5 w Google.")

    save(doc, os.path.join(OUTDIR, "kawiarnia.docx"))


# ════════════════════════════════════════════════════════════════════════
# WARSZTATY
# ════════════════════════════════════════════════════════════════════════
def gen_warsztaty():
    reset_alt()
    doc, t = new_doc("Warsztaty")

    sec(t, "1. HERO — banner powitalny")
    row(t, "BREADCRUMB", "Strona główna › Warsztaty")
    row(t, "NAGŁÓWEK H1 (linia 1)", "Warsztaty dla dzieci.", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 2 — kursywa)", "Nowy Targ — plastyka, glina,", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 3)", "joga, animaloterapia.", bold=True, size=11)
    row(t, "PRZYCISK 1", "Sprawdź terminy na Instagramie")
    row(t, "PRZYCISK 2", "Zadzwoń · 693 766 049")

    sec(t, "2. INTRO — direct answer (pierwsze 100 słów)")
    row(t, "AKAPIT INTRO", "Warsztaty dla dzieci w Bawiszu w Nowym Targu prowadzą zaproszeni partnerzy — w drewnianej sali Montessori przy ul. Krzywej 19B. Repertuar zmienia się: plastyka, glina, joga dla dzieci, animaloterapia, sensoryka. Wszystkie najbliższe terminy ogłaszamy na Instagramie @bawisz_bawialnia — tam znajdziesz datę, prowadzącego, wiek dziecka i cenę. Standardowo 1,5 h zajęć, 60-80 zł od osoby.")

    sec(t, "3. PRZYKŁADY — 5 warsztatów")
    row(t, "EYEBROW", "[ Co już u nas było ]")
    row(t, "NAGŁÓWEK H2", "Przykłady warsztatów.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Warsztaty robimy nieregularnie — z różnymi prowadzącymi i tematami. Poniżej kilka, które już u nas były. Najbliższe terminy zawsze na Instagramie.")
    warsztaty = [
        ('Akademia Maluszka', 'do 4 roku życia', '1,5 h · 60 zł/os.', 'Anna Jagoda Mrożek · Effata Anima',
         'Cykl spotkań plastycznych dla najmłodszych — zabawa kolorem, proste prace plastyczne, swobodna zabawa po zajęciach. Cztery terminy w sezonie.'),
        ('Mały Artysta', 'od 4 roku życia', '1,5 h · 70 zł/os.', 'Anna Jagoda Mrożek · Effata Anima',
         'Cykl plastyczny „Poznajemy mistrzów, tworzymy po swojemu". Dzieci poznają wybranego twórcę i robią własną pracę inspirowaną jego stylem.'),
        ('Warsztaty z gliny', 'od 4 roku życia', '1,5 h · 80 zł/os.', 'Anna Jagoda Mrożek · Effata Anima',
         'Lepienie z gliny — od ugniatania i poznawania faktury, po formowanie pierwszych figurek. Pracę dziecko zabiera do domu po wyschnięciu.'),
        ('Gimnastyka dla Smyka', 'od 2 roku życia', '1,5 h · 60 zł/os.', 'Anna Tomalak · Asana Joga',
         'Cykliczne zajęcia ruchowe dla najmłodszych — joga dla dzieci, pozycje zwierząt, zabawy z chustą i piłką. Bardziej zabawa niż „trening".'),
        ('Warsztaty z kurami jedwabistymi', 'dla dzieci i rodzin', 'animaloterapia · sensoryka', 'partner zewnętrzny',
         'Spotkanie z puszystymi kurami jedwabistymi — przytulanie, karmienie, dotykanie miękkiego pierza. Element wyciszenia, kontaktu z naturą i sensoryki.'),
    ]
    for i, (h, age, meta, by, p) in enumerate(warsztaty, 1):
        row(t, f"WARSZTAT {i} — tytuł", h, bold=True)
        row(t, f"WARSZTAT {i} — wiek", age)
        row(t, f"WARSZTAT {i} — czas/cena", meta)
        row(t, f"WARSZTAT {i} — prowadzący", by)
        row(t, f"WARSZTAT {i} — opis", p)

    sec(t, "4. GALERIA — podpisy zdjęć")
    row(t, "EYEBROW", "[ Galeria ]")
    row(t, "NAGŁÓWEK H2", "Tak wyglądają warsztaty w Bawiszu.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Plastyka, glina, sensoryka i joga w drewnianej sali. Kliknij zdjęcie, żeby powiększyć.")
    row(t, "ALT zdjęcia 1", "Warsztaty plastyczne w Bawiszu")
    row(t, "ALT zdjęcia 2", "Dzieci na warsztacie")
    row(t, "ALT zdjęcia 3", "Sensoplastyka")
    row(t, "ALT zdjęcia 4", "Glina dla dzieci")
    row(t, "ALT zdjęcia 5", "Joga dla dzieci")

    sec(t, "5. PROCES — 4 kroki")
    row(t, "EYEBROW", "[ Jak to wygląda ]")
    row(t, "NAGŁÓWEK H2", "Cztery kroki. Zaczynamy od Instagrama.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Najczęstsze pytanie: „skąd mam wiedzieć, kiedy będzie warsztat?\". Odpowiedź: ogłaszamy je z 1-2 tygodniowym wyprzedzeniem na Instagramie — zaobserwuj profil, żeby nie przegapić.")
    row(t, "KROK 1 — tytuł", "Śledź Instagram", bold=True)
    row(t, "KROK 1 — opis", "Najbliższe warsztaty ogłaszamy na profilu @bawisz_bawialnia z 1-2 tygodniowym wyprzedzeniem. W ogłoszeniu znajdziesz datę, prowadzącego, dla jakiego wieku, ile trwa i ile kosztuje. Włącz powiadomienia — część terminów wyprzedaje się w kilka dni.")
    row(t, "KROK 2 — tytuł", "Zadzwoń, żeby zarezerwować", bold=True)
    row(t, "KROK 2 — opis", "Po zobaczeniu ogłoszenia dzwonisz na +48 693 766 049 albo piszesz na Instagramie. Mówisz, na który warsztat i ile dzieci. Potwierdzamy miejsce — liczba miejsc ograniczona, zwykle 6-10 dzieci na warsztat.")
    row(t, "KROK 3 — tytuł", "Zajęcia z prowadzącą", bold=True)
    row(t, "KROK 3 — opis", "1,5 godziny zajęć z zaproszoną prowadzącą — pokazuje technikę, prowadzi zabawę, pomaga przy trudniejszych krokach. Zostajesz w sali, jeśli chcesz, albo czekasz w kawiarni obok. Przy najmłodszych warto być na sali.")
    row(t, "KROK 4 — tytuł", "Po zajęciach — kawa i ciasto", bold=True)
    row(t, "KROK 4 — opis", "Dziecko zabiera swoją pracę do domu (jeśli warsztat plastyczny). Wy macie chwilę na kawę i ciasto domowe w kawiarni przy sali. Materiały sprząta prowadząca — wracacie spokojnie do auta.")

    sec(t, "6. MID-CTA — środkowy banner")
    row(t, "NAGŁÓWEK H2", "Najbliższy warsztat?", bold=True, size=11)
    row(t, "OPIS", "Aktualne terminy, ceny i zapisy na profilu @bawisz_bawialnia. Jeśli masz pomysł na warsztat tematyczny dla zamkniętej grupy (urodziny, wyjście przedszkolne) — dzwoń, ustalamy minimum 2 tygodnie wcześniej.")
    row(t, "PRZYCISK 1", "Otwórz Instagram")
    row(t, "PRZYCISK 2", "Zadzwoń · 693 766 049")

    sec(t, "7. FAQ — najczęstsze pytania")
    warsztaty_faq = [
        ('Skąd dowiem się o najbliższych warsztatach w Bawiszu?',
         'Wszystkie warsztaty ogłaszamy na bieżąco na Instagramie @bawisz_bawialnia — tam znajdziesz datę, godzinę, prowadzącego, wiek dziecka i cenę. Najszybciej zaobserwować profil i włączyć powiadomienia, bo część terminów wyprzedaje się w 2-3 dni.'),
        ('Kto prowadzi warsztaty?',
         'Zapraszamy zewnętrznych prowadzących, którzy specjalizują się w swoim temacie. Stale pracujemy m.in. z Anną Jagodą Mrożek (Effata Anima — Akademia Maluszka, Mały Artysta, warsztaty z gliny) oraz Anną Tomalak (Asana Joga — Gimnastyka dla Smyka). Co jakiś czas dochodzą jednorazowi partnerzy do warsztatów tematycznych (np. animaloterapia z kurami jedwabistymi).'),
        ('Jakie warsztaty były dotychczas u was?',
         'Cykliczne: Akademia Maluszka (do 4 lat, plastyka), Mały Artysta (od 4 lat, plastyka), Gimnastyka dla Smyka (od 2 lat, ruch). Jednorazowe: warsztaty z gliny, warsztaty z kurami jedwabistymi (animaloterapia), zajęcia tematyczne wokół pór roku i świąt. Repertuar zmienia się — sprawdź Instagram, co planujemy w najbliższym miesiącu.'),
        ('Ile kosztują warsztaty i jak długo trwają?',
         'Standardowo 1,5 godziny zajęć z prowadzącą plus swobodna zabawa. Cena 60-80 zł od osoby — dokładna kwota podana jest zawsze przy ogłoszeniu konkretnego warsztatu na Instagramie. Cena obejmuje materiały (farby, glinę, masy sensoryczne i to, co potrzebne na danym warsztacie).'),
        ('Dla jakiego wieku dziecka są warsztaty?',
         'Każdy warsztat ma określony wiek — od 2 lat na Gimnastyce dla Smyka, do 4 lat na Akademii Maluszka, od 4 lat na Małym Artyście i warsztatach z gliny, dla dzieci i rodzin razem na warsztatach z kurami. Wiek jest zawsze podany w ogłoszeniu na Instagramie.'),
        ('Jak się zapisać?',
         'Zapisy przez telefon: +48 693 766 049 albo wiadomość na Instagramie. Liczba miejsc jest ograniczona (zwykle 6-10 dzieci na warsztat), więc dzwoń od razu po zobaczeniu ogłoszenia. Po telefonie potwierdzamy rezerwację i przypominamy o terminie dzień wcześniej.'),
        ('Mam pomysł na warsztat — czy mogę go u was poprowadzić?',
         'Tak — szukamy nowych partnerów. Prowadzisz zajęcia dla dzieci (plastyka, muzyka, ruch, zwierzęta, kuchnia), masz portfolio i chcesz wynająć drewnianą salę Montessori w Nowym Targu? Napisz na Instagramie albo zadzwoń na +48 693 766 049 — ustalimy szczegóły.'),
    ]
    row(t, "EYEBROW", "[ FAQ ]")
    row(t, "NAGŁÓWEK H2", "Najczęstsze pytania.", bold=True, size=11)
    for i, (q, a) in enumerate(warsztaty_faq, 1):
        row(t, f"PYTANIE {i}", q, bold=True)
        row(t, f"ODPOWIEDŹ {i}", a)

    sec(t, "8. FINAL CTA — końcowy banner")
    row(t, "NAGŁÓWEK H2", "Zapisy na warsztaty", bold=True, size=11)
    row(t, "OPIS", "ul. Krzywa 19B, Nowy Targ. Aktualny grafik warsztatów (plastyka, glina, joga, animaloterapia) na Instagramie @bawisz_bawialnia. Zapisy telefonicznie albo wiadomością na Instagramie.")
    row(t, "PRZYCISK 1", "Otwórz Instagram")
    row(t, "PRZYCISK 2", "Zadzwoń · 693 766 049")

    sec(t, "9. META — SEO")
    row(t, "META TITLE", "Warsztaty dla dzieci Nowy Targ — Bawisz | plastyka, glina, joga")
    row(t, "META DESCRIPTION", "Warsztaty dla dzieci w Nowym Targu w drewnianej sali Montessori: plastyka, glina, joga dla dzieci, animaloterapia. Terminy na Instagramie. 60-80 zł/os, 1,5 h.")

    save(doc, os.path.join(OUTDIR, "warsztaty.docx"))


# ════════════════════════════════════════════════════════════════════════
# O NAS
# ════════════════════════════════════════════════════════════════════════
def gen_o_nas():
    reset_alt()
    doc, t = new_doc("O nas")

    sec(t, "1. HERO — banner powitalny")
    row(t, "BREADCRUMB", "Strona główna › O nas")
    row(t, "NAGŁÓWEK H1 (linia 1)", "Bawialnia Montessori", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 2 — kursywa)", "w Nowym Targu.", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 3)", "Drewniana, naturalna, dla dzieci 0-10 lat.", bold=True, size=11)
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "@bawisz_bawialnia")

    sec(t, "2. INTRO — direct answer (pierwsze 100 słów)")
    row(t, "AKAPIT INTRO", "Bawisz to bawialnia Montessori i kawiarnia pod jednym dachem przy ul. Krzywej 19B w Nowym Targu. Pedagogika Montessori w praktyce — drewniane wnętrze, naturalne zabawki z drewna i sklejki, samodzielność dziecka. Bez plastiku, bez krzykliwego oświetlenia, bez animatorów zabawiających dzieci za rodzica. Dzieci od 0 do 10 lat bawią się ze swoimi rodzicami, ty pijesz kawę specialty obok. Wstęp do bawialni od 25 zł za godzinę, kawiarnia bez biletu. Ocena 4.9/5 w Google.")

    sec(t, "3. WARTOŚCI — 4 cards (Filozofia)")
    row(t, "EYEBROW", "[ Filozofia ]")
    row(t, "NAGŁÓWEK H2", "Cztery rzeczy, które robimy inaczej.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Najczęstsze pytanie po pierwszej wizycie: „dlaczego u was jest tak spokojnie?\". Odpowiedź jest w czterech decyzjach, które podjęliśmy zanim wstawiliśmy pierwszą półkę.")
    wartosci = [
        ('1', 'Drewno, nie plastik.',
         'Cała architektura sali jest z drewna i sklejki. Zabawki są z drewna, sklejki i tkanin — bez krzykliwych kolorów, bez plastikowego hałasu, bez baterii. To nie estetyka, tylko spokojna przestrzeń, w której dziecku łatwiej się skupić, a ty słyszysz własne myśli.'),
        ('2', 'Wszystko na wysokości dziecka.',
         'Półki niskie, koszyki dostępne, każda zabawka w swoim miejscu. Dziecko samo wybiera, co chce robić, samo odkłada na miejsce. To jest sedno Montessori — samodzielność, nie wykonywany program zajęć ani zadania od kogoś dorosłego.'),
        ('3', 'Bawisz się ze swoim dzieckiem.',
         'Nie zabieramy ci dziecka. Nie ma animatorów, nie ma „zostaw je u nas, my się zajmiemy". Bawicie się razem — przy klockach, sensoplastyce, w kąciku czytelniczym, w strefie zabaw w role. Sala jest tak przemyślana, że nie musisz pilnować na każdym kroku.'),
        ('4', 'Kawiarnia obok, nie poczekalnia.',
         'Po drugiej stronie szyby parzymy specialty espresso, flat white, latte i matcha latte. Pieczemy ciasta domowe — sezonowe, brownie, sernik, szarlotka, w wersjach bezglutenowych i bez cukru. Dla dzieci owoce, mleko, soczki bez cukru, kanapka z dżemem. Kawa zostaje gorąca, dziecko ma wszystko pod ręką.'),
    ]
    for n, h, p in wartosci:
        row(t, f"WARTOŚĆ {n} — tytuł", h, bold=True)
        row(t, f"WARTOŚĆ {n} — opis", p)

    sec(t, "4. MID-CTA — środkowy banner")
    row(t, "NAGŁÓWEK H2", "Wpadnij na kawę.", bold=True, size=11)
    row(t, "OPIS", "Bez rezerwacji w tygodniu — wchodzisz, kupujesz bilet do bawialni przy ladzie (25 zł za godzinę), bierzesz kawę z lady. Urodziny, warsztaty i grupy zorganizowane ustalamy przez telefon albo Instagram.")
    row(t, "PRZYCISK 1", "693 766 049")
    row(t, "PRZYCISK 2", "Napisz na Instagramie")

    sec(t, "5. FAQ — najczęstsze pytania")
    onas_faq = [
        ('Czy w Bawiszu są animatorzy, którzy zajmą się dzieckiem?',
         'Nie. W codziennym wstępie do bawialni nie ma animatorów. Bawisz się ze swoim dzieckiem przy naturalnych zabawkach z drewna, my zajmujemy się salą, kawiarnią i wszystkim, co dookoła. Animatorzy i prowadzący pojawiają się tylko na zorganizowanych urodzinach (pakiet MINI 45 zł/os. lub STANDARD 74 zł/os.) i na warsztatach.'),
        ('Co właściwie znaczy „bawialnia Montessori" w praktyce?',
         'W praktyce to drewniana sala, w której wszystko jest na wysokości dziecka — półki, koszyki z zabawkami, kąciki tematyczne. Zabawki są z drewna, sklejki i tkanin, bez plastiku i bez baterii. Dziecko samo wybiera, co robi, samo odkłada na miejsce. To filozofia samodzielności i naturalnych materiałów, nie marka mebli ani sztywny program zajęć.'),
        ('W jakim wieku dzieci mogą tu przyjść?',
         'Sala jest dla dzieci od 0 do 10 lat. Wewnątrz dzieli się na strefy dopasowane do wieku — najmłodsze (do 2 lat) bawią się w spokojniejszej części z miękkim podłożem, 3-5-latki w strefie z klockami i zabawami w role, starsze (6-10 lat) w kącikach kreatywnych. Opiekun zawsze gratis, drugie dziecko w rodzinie 25% taniej.'),
        ('Czy do kawiarni potrzebny jest bilet do bawialni?',
         'Nie. Kawiarnia działa osobno. Wpadasz na kawę bez dziecka albo z dzieckiem, które tego dnia nie chce się bawić — wchodzisz, zamawiasz przy ladzie, siadasz przy stoliku. Bilet do bawialni kupujesz tylko wtedy, gdy dziecko ma wejść do sali (25 zł za 1 h, 33 zł za 2 h, 45 zł cały dzień, 215 zł karnet miesięczny no limit).'),
        ('Gdzie jesteście i czy jest parking?',
         'ul. Krzywa 19B, 34-400 Nowy Targ — centrum, dwie minuty od Rynku. Parking jest tuż obok lokalu, bez problemu zaparkujesz nawet w sobotę po południu. Otwarte codziennie od 10:00 (w tygodniu do 19:00, w weekendy do 20:00).'),
        ('Czy organizujecie urodziny, warsztaty i wyjścia grupowe?',
         'Tak — w trzech formatach. Urodziny w pakiecie MINI (45 zł/os., 2 h) lub STANDARD (74 zł/os., 2,5 h, sala tylko dla was, min. 10 dzieci, prezent dla solenizanta). Warsztaty — sensoplastyka, plastyka, zajęcia kreatywne w duchu Montessori. Wyjścia grupowe dla przedszkoli i szkół od 15 zł za dziecko (min. 10 osób, kawa lub herbata gratis dla 1 opiekuna na 5 dzieci). Detale na osobnych podstronach.'),
        ('Jak najszybciej zarezerwować termin albo zapytać o szczegóły?',
         'Telefon lub SMS na +48 693 766 049 — odpisujemy w godzinach otwarcia. Można też napisać na Instagramie (@bawisz_bawialnia) albo na Facebooku. Wstęp do bawialni jest bez rezerwacji — wchodzisz, kupujesz bilet przy ladzie. Urodziny, warsztaty i grupy zorganizowane wymagają wcześniejszego ustalenia daty.'),
    ]
    row(t, "EYEBROW", "[ FAQ ]")
    row(t, "NAGŁÓWEK H2", "Najczęstsze pytania.", bold=True, size=11)
    for i, (q, a) in enumerate(onas_faq, 1):
        row(t, f"PYTANIE {i}", q, bold=True)
        row(t, f"ODPOWIEDŹ {i}", a)

    sec(t, "6. FINAL CTA — końcowy banner")
    row(t, "NAGŁÓWEK H2", "Wpadnij do nas.", bold=True, size=11)
    row(t, "OPIS", "Sala otwarta od 10:00, kawa parzona od pierwszego klienta. Wstęp do bawialni od 25 zł za godzinę, opiekun zawsze gratis. ul. Krzywa 19B w Nowym Targu — centrum, parking obok.")
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "Zarezerwuj urodziny")

    sec(t, "7. META — SEO")
    row(t, "META TITLE", "Bawialnia Montessori Nowy Targ | drewniana sala — Bawisz")
    row(t, "META DESCRIPTION", "Bawialnia Montessori w Nowym Targu — drewniana sala, naturalne zabawki, kawiarnia obok. Dzieci 0-10 lat. Wstęp od 25 zł. ul. Krzywa 19B. Ocena 4.9/5.")

    save(doc, os.path.join(OUTDIR, "o-nas.docx"))


# ════════════════════════════════════════════════════════════════════════
# OFERTA GRUPOWA (Dla przedszkoli)
# ════════════════════════════════════════════════════════════════════════
def gen_oferta_grupowa():
    reset_alt()
    doc, t = new_doc("Dla przedszkoli (oferta grupowa)")

    sec(t, "1. HERO — banner powitalny")
    row(t, "BREADCRUMB", "Strona główna › Dla przedszkoli")
    row(t, "NAGŁÓWEK H1 (linia 1)", "Oferta dla przedszkoli.", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 2 — kursywa)", "Nowy Targ — drewniana", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 3)", "bawialnia Montessori.", bold=True, size=11)
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "Napisz na Instagramie")

    sec(t, "2. INTRO — direct answer (pierwsze 100 słów)")
    row(t, "AKAPIT INTRO", "Wyjścia grupowe dla przedszkoli i szkół z Nowego Targu i okolic w drewnianej bawialni Montessori przy ul. Krzywej 19B. Minimum 10 dzieci, od 15 zł za godzinę za dziecko (23 zł za 2 godziny, 35 zł NO LIMIT na cały dzień). Kawa lub herbata gratis dla 1 opiekuna na 5 dzieci. Sala dla wieku 0-10 lat, naturalne zabawki, strefy dopasowane do wieku. Faktura dla placówki z odroczonym terminem płatności. Termin rezerwujemy telefonicznie — najlepiej z 1-2 tygodniami wyprzedzenia.")

    sec(t, "3. DLA KOGO — 3 typy grup")
    row(t, "EYEBROW", "[ Dla kogo ]")
    row(t, "NAGŁÓWEK H2", "Trzy rodzaje wyjść grupowych.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Cennik jest jeden, sala ta sama — drewniana, podzielona na strefy według wieku. Rodzaj wyjścia ma znaczenie po to, żebyśmy ustawili strefy pod konkretne grupy wiekowe i przygotowali poczęstunek pod dietę grupy.")
    audiences = [
        ('Przedszkola', 'Wyjścia dla przedszkoli',
         'Cała grupa wchodzi razem, dzieci bawią się przy drewnianych zabawkach Montessori pod opieką swoich nauczycieli. Strefy są podzielone według wieku (3-4 lata, 5-6 lat), więc 6-latki nie wchodzą w paradę 3-latkom. Idealne na wycieczkę pieszą z budynku przedszkola — jesteśmy 5-7 minut spacerem od centrum Nowego Targu.'),
        ('Szkoły podstawowe', 'Bawialnia dla szkół',
         'Klasy 0-3 mieszczą się komfortowo w drewnianej sali. Świetne na zakończenie roku, dzień dziecka, mikołajki klasowe albo jako nagroda za projekt. Dzień otwarty (NO LIMIT) sprawdza się przy klasach łączonych — dzieci same wybierają, gdzie chcą się bawić, nauczyciel ma je w polu widzenia.'),
        ('Grupy zorganizowane', 'Imprezy dla grup',
         'Spotkania urodzinowe całych klas, zajęcia integracyjne dla zespołów dziecięcych, wyjścia takich grup jak klub mam, drużyna harcerska albo świetlica. Łączymy wejście grupowe z dodatkami: poczęstunek, warsztat sensoplastyki, dekoracje urodzinowe — co potrzeba, to dokładamy.'),
    ]
    for i, (kategoria, h, p) in enumerate(audiences, 1):
        row(t, f"GRUPA {i} — kategoria", kategoria, bold=True)
        row(t, f"GRUPA {i} — tytuł", h, bold=True)
        row(t, f"GRUPA {i} — opis", p)

    sec(t, "4. PAKIETY CZASOWE — 3 ceny")
    row(t, "EYEBROW", "[ Pakiety czasowe ]")
    row(t, "NAGŁÓWEK H2", "Trzy ceny. Bez gwiazdek.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Cena za dziecko za czas zabawy. Opiekunowie wchodzą gratis, 1 na 5 dzieci ma kawę albo herbatę z naszej kawiarni w cenie pakietu.")
    pakiety = [
        ('1 godzina', '15 zł', '/ dziecko', None,
         'Krótkie wyjście, pełen dostęp do drewnianej sali. Sprawdza się, gdy grupa wraca do przedszkola na obiad albo łączy wizytę z innym punktem programu.'),
        ('2 godziny', '23 zł', '/ dziecko', 'Najczęstszy wybór',
         'Czas, by każde dziecko spróbowało wszystkich stref: zabawy w role, kącik czytelniczy, drewniane zabawki sensoryczne, sala ruchowa. Najczęstszy wybór przedszkoli z Nowego Targu i okolic.'),
        ('NO LIMIT', '35 zł', '/ dziecko · cały dzień', None,
         'Cały dzień otwarcia — od 10:00 do zamknięcia. Kawa i herbata bez limitu dla opiekunów (1:5). Dobre na wycieczki całodniowe i dni, kiedy dzieci wracają do sali po przerwie obiadowej.'),
    ]
    for i, (czas, cena, jednostka, badge, p) in enumerate(pakiety, 1):
        row(t, f"PAKIET {i} — czas", czas, bold=True)
        row(t, f"PAKIET {i} — cena", cena)
        row(t, f"PAKIET {i} — jednostka", jednostka)
        if badge:
            row(t, f"PAKIET {i} — badge", badge)
        row(t, f"PAKIET {i} — opis", p)
    row(t, "NOTKA POD PAKIETAMI", "Minimum: 10 dzieci w grupie. Kawa lub herbata gratis dla 1 opiekuna na 5 dzieci. Poczęstunek dla dzieci na zamówienie — ustalamy zakres telefonicznie.")

    sec(t, "5. GALERIA — podpisy zdjęć")
    row(t, "EYEBROW", "[ Galeria ]")
    row(t, "NAGŁÓWEK H2", "Tak wyglądają wyjścia grupowe.", bold=True, size=11)
    row(t, "PODTYTUŁ", "220 m² drewna, strefy podzielone wiekiem, opiekun w polu widzenia. Kliknij zdjęcie, żeby powiększyć.")
    row(t, "ALT zdjęcia 1", "Grupa przedszkolna w Bawiszu")
    row(t, "ALT zdjęcia 2", "Sala dla grup")
    row(t, "ALT zdjęcia 3", "Wyjście grupowe")
    row(t, "ALT zdjęcia 4", "Strefy zabaw dla grup")
    row(t, "ALT zdjęcia 5", "Dzieci w sali Montessori")

    sec(t, "6. PROCES — 4 kroki")
    row(t, "EYEBROW", "[ Jak to wygląda ]")
    row(t, "NAGŁÓWEK H2", "Cztery kroki. Od telefonu do faktury.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Najczęstsze pytanie: „co dokładnie robicie wy, a co mam zrobić ja?\". Odpowiedź: wy przyprowadzacie grupę, my robimy resztę — od przygotowania sali po fakturę.")
    proces = [
        ('1', 'Telefon i ustalenie szczegółów',
         'Dzwonisz na +48 693 766 049 albo piszesz na Instagramie. Podajesz datę, liczbę dzieci, czas trwania (1 h, 2 h albo NO LIMIT) oraz to, czy chcecie poczęstunek. Sprawdzamy wolny termin i potwierdzamy rezerwację.'),
        ('2', 'Przygotowanie sali',
         'Zanim grupa przyjdzie, ustawiamy strefy według wieku dzieci, sprawdzamy bezpieczeństwo, parzymy kawę i herbatę dla opiekunów. Jeśli zamówiliście poczęstunek — przygotowujemy go tak, żeby był gotowy na wejście grupy.'),
        ('3', 'Wejście grupy i zabawa',
         'Dzieci wchodzą i bawią się przy drewnianych zabawkach Montessori — sensoplastyka, zabawy w role, kącik czytelniczy, sala ruchowa. Nauczyciele zostają z grupą, my dbamy o przestrzeń, poczęstunek i to, żeby nic nie zabrakło.'),
        ('4', 'Rozliczenie z placówką',
         'Wystawiamy fakturę dla przedszkola lub szkoły z odroczonym terminem płatności (przelew po wizycie). Potrzebujemy NIP placówki i danych do faktury — przesyłacie je SMS-em albo mailem po ustaleniu terminu.'),
    ]
    for n, h, p in proces:
        row(t, f"KROK {n} — tytuł", h, bold=True)
        row(t, f"KROK {n} — opis", p)

    sec(t, "7. MID-CTA — środkowy banner")
    row(t, "NAGŁÓWEK H2", "Termin wyjścia?", bold=True, size=11)
    row(t, "OPIS", "Wolne terminy w godzinach porannych i wczesnym popołudniem (od 10:00 do około 14:00 — przed otwarciem dla rodzin indywidualnych). Najszybciej przez telefon — od razu sprawdzamy wolne daty i ustalamy szczegóły.")
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "Napisz na Instagramie")

    sec(t, "8. FAQ — najczęstsze pytania")
    og_faq = [
        ('Ile kosztuje wyjście grupowe dla przedszkola w Bawiszu?',
         'Cennik startuje od 15 zł za dziecko za godzinę zabawy. 2 godziny — 23 zł, cały dzień (NO LIMIT) — 35 zł. Minimum to grupa od 10 dzieci. Kawa lub herbata dla opiekunów gratis (1 opiekun na 5 dzieci).'),
        ('Od ilu dzieci minimum mogę zarezerwować wyjście grupowe?',
         'Minimum to 10 dzieci. Mniejsze grupy obsługujemy po cenie indywidualnego biletu wstępu (od 25 zł za godzinę za dziecko). Górnej granicy nie podajemy z głowy — przy większej grupie ustalamy szczegóły organizacyjne telefonicznie, żeby każde dziecko miało komfort.'),
        ('Czy opiekunowie wchodzą gratis i dostają coś do picia?',
         'Tak. Opiekunowie wchodzą bez opłaty, a 1 opiekun na 5 dzieci dostaje kawę albo herbatę z naszej kawiarni gratis. Pozostali opiekunowie zamawiają z menu na miejscu — mamy kawę specialty, ciasta domowe i napoje dla dzieci.'),
        ('Czy zapewniacie poczęstunek dla dzieci?',
         'Tak, na zamówienie. Standardowo poczęstunku w cenie wejścia nie ma — możemy go przygotować dodatkowo, ustalamy zakres (np. paluszki, owoce, soczki, kanapki) telefonicznie. Cena dopasowana do liczby dzieci i tego, co dokładnie zamawiacie.'),
        ('Z jakim wyprzedzeniem trzeba rezerwować termin?',
         'Najlepiej 1-2 tygodnie wcześniej. W sezonie (październik-grudzień, marzec-czerwiec) terminy zapełniają się szybko, więc warto pisać miesiąc wcześniej. W okresach spokojniejszych udaje się zarezerwować nawet w tym samym tygodniu — najszybciej przez telefon: +48 693 766 049.'),
        ('Czy wystawiacie fakturę dla placówki?',
         'Tak. Wystawiamy fakturę dla przedszkola lub szkoły z odroczonym terminem płatności (przelew po wizycie). Potrzebujemy NIP placówki i danych do faktury — najlepiej wysłać je SMS-em albo na Instagramie po ustaleniu terminu.'),
        ('Dla jakiego wieku dzieci jest sala i czy mamy wyłączność?',
         'Sala jest dla dzieci od 0 do 10 lat — strefy są podzielone według wieku, więc 3-latek bawi się w innej części niż 7-latek. Wyłączność na sali nie wchodzi w cenę pakietu grupowego, ale w godzinach porannych (przed 12:00 w dni powszednie) zwykle jesteście jedyną grupą. Pełną wyłączność można wykupić indywidualnie — pytaj telefonicznie.'),
        ('Czy można połączyć wyjście grupowe z urodzinami przedszkolaka?',
         'Tak — robimy mieszane wyjścia: rano wejście grupowe dla całej grupy, później pakiet urodzinowy dla solenizanta. Łączymy oba cenniki, ustalamy szczegóły i poczęstunek telefonicznie. Najprościej napisać, ile dzieci, jakie urodziny i jaki czas trwania.'),
    ]
    row(t, "EYEBROW", "[ FAQ ]")
    row(t, "NAGŁÓWEK H2", "Najczęstsze pytania.", bold=True, size=11)
    for i, (q, a) in enumerate(og_faq, 1):
        row(t, f"PYTANIE {i}", q, bold=True)
        row(t, f"ODPOWIEDŹ {i}", a)

    sec(t, "9. FINAL CTA — końcowy banner")
    row(t, "NAGŁÓWEK H2", "Rezerwacja wyjścia grupowego", bold=True, size=11)
    row(t, "OPIS", "ul. Krzywa 19B, Nowy Targ. Pakiety od 15 zł za godzinę za dziecko, faktura dla placówki z odroczonym terminem płatności. Termin i szczegóły ustalamy telefonicznie — najlepiej 1-2 tygodnie wcześniej.")
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "Napisz na Instagramie")

    sec(t, "10. META — SEO")
    row(t, "META TITLE", "Oferta dla przedszkoli Nowy Targ — Bawisz | od 15 zł/dziecko")
    row(t, "META DESCRIPTION", "Wyjścia grupowe dla przedszkoli i szkół w Nowym Targu — drewniana bawialnia Montessori. Od 15 zł za godzinę za dziecko, kawa gratis dla opiekunów. Min. 10 dzieci.")

    save(doc, os.path.join(OUTDIR, "oferta-grupowa.docx"))


# ════════════════════════════════════════════════════════════════════════
# KONTAKT
# ════════════════════════════════════════════════════════════════════════
def gen_kontakt():
    reset_alt()
    doc, t = new_doc("Kontakt")

    sec(t, "1. HERO — banner powitalny")
    row(t, "BREADCRUMB", "Strona główna › Kontakt")
    row(t, "NAGŁÓWEK H1 (linia 1)", "Kontakt — Bawisz.", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 2 — kursywa)", "Nowy Targ, Krzywa 19B —", bold=True, size=11)
    row(t, "NAGŁÓWEK H1 (linia 3)", "telefon, Instagram, mapa.", bold=True, size=11)
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "@bawisz_bawialnia")

    sec(t, "2. INTRO — direct answer (pierwsze 100 słów)")
    row(t, "AKAPIT INTRO", "Najszybszy kontakt z Bawiszem w Nowym Targu to telefon lub SMS na +48 693 766 049 — odbieramy codziennie od 10:00 (w tygodniu do 19:00, w weekendy do 20:00). Możesz też napisać na Instagramie albo na Messengerze, albo po prostu wpaść — ul. Krzywa 19B, 34-400 Nowy Targ, dwie minuty od Rynku, parking obok lokalu. Wstęp do bawialni jest bez rezerwacji. Urodziny, warsztaty i wyjścia grupowe ustalamy wcześniej, najlepiej telefonem.")

    sec(t, "3. KANAŁY KONTAKTU — 4 kanały")
    row(t, "EYEBROW", "[ Cztery sposoby ]")
    row(t, "NAGŁÓWEK H2", "Wybierz, jak się łapiemy.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Każdy kanał obsługujemy w godzinach otwarcia. Telefon i SMS to najszybsza droga, Instagram i Messenger sprawdzamy w tym samym rytmie.")
    channels = [
        ('+48 693 766 049', 'Odbieramy najszybciej. SMS poza godzinami — oddzwaniamy rano.'),
        ('@bawisz_bawialnia', 'Codzienne kadry z sali, stories z urodzin. DM jak telefon.'),
        ('Bawisz · bawialnia Montessori', 'Wydarzenia, warsztaty, info dla grup. Messenger sprawdzamy w godzinach.'),
        ('Krzywa 19B, Nowy Targ', 'Centrum miasta, dwie minuty od Rynku, parking obok lokalu.'),
    ]
    for i, (label, sub) in enumerate(channels, 1):
        row(t, f"KANAŁ {i} — etykieta", label, bold=True)
        row(t, f"KANAŁ {i} — opis", sub)

    sec(t, "4. ADRES + GODZINY + MAPA")
    row(t, "EYEBROW", "[ Godziny i adres ]")
    row(t, "NAGŁÓWEK H2", "Codziennie otwarte.", bold=True, size=11)
    row(t, "PODTYTUŁ", "Otwarte przez cały rok — nie zamykamy się w wakacje ani między świętami. Ostatnie wejście do bawialni przyjmujemy 30 minut przed zamknięciem.")
    row(t, "GODZINY — Pn", "Poniedziałek: 10:00 — 19:00")
    row(t, "GODZINY — Wt", "Wtorek: 10:00 — 19:00")
    row(t, "GODZINY — Śr", "Środa: 10:00 — 19:00")
    row(t, "GODZINY — Cz", "Czwartek: 10:00 — 19:00")
    row(t, "GODZINY — Pt", "Piątek: 10:00 — 19:00")
    row(t, "GODZINY — Sob", "Sobota: 10:00 — 20:00 (dłużej)")
    row(t, "GODZINY — Nd", "Niedziela: 10:00 — 20:00 (dłużej)")
    row(t, "KARTA ADRES — eyebrow", "[ Jak do nas trafić ]")
    row(t, "KARTA ADRES — adres", "ul. Krzywa 19B, Nowy Targ", bold=True)
    row(t, "ADRES — kod + miasto", "34-400 Nowy Targ", bold=True)
    row(t, "ADRES — opis", "Centrum, parking obok lokalu")
    row(t, "TELEFON", "+48 693 766 049", bold=True)
    row(t, "TELEFON — podpis", "Telefon lub SMS — w godzinach otwarcia")
    row(t, "GODZINY — tytuł", "Codziennie od 10:00", bold=True)
    row(t, "GODZINY — podpis", "Pn–Pt do 19:00, Sb–Nd do 20:00")
    row(t, "PRZYCISK CTA", "Nawiguj do nas")

    sec(t, "5. MID-CTA — środkowy banner")
    row(t, "NAGŁÓWEK H2", "Termin urodzin?", bold=True, size=11)
    row(t, "OPIS", "Pakiet STANDARD rezerwuje się 3-6 tygodni wcześniej, MINI często łapiemy w tym samym tygodniu. Najszybciej telefonem — od razu sprawdzamy wolne daty i ustalamy motyw dekoracji. Prowadząca poprowadzi program — ty siedzisz przy kawie.")
    row(t, "PRZYCISK 1", "Zarezerwuj urodziny")
    row(t, "PRZYCISK 2", "Zadzwoń · 693 766 049")

    sec(t, "6. FAQ — najczęstsze pytania")
    kontakt_faq = [
        ('Jak najszybciej skontaktować się z Bawiszem w Nowym Targu?',
         'Telefon lub SMS na +48 693 766 049 — odbieramy w godzinach otwarcia, czyli codziennie od 10:00 (w tygodniu do 19:00, w weekendy do 20:00). Poza godzinami zostaw SMS, oddzwaniamy rano. Można też napisać na Instagramie (@bawisz_bawialnia) albo na Facebooku — Messengera sprawdzamy w tym samym rytmie co telefon.'),
        ('Gdzie jest Bawisz i czy łatwo zaparkować?',
         'ul. Krzywa 19B, 34-400 Nowy Targ — centrum, dwie minuty od Rynku. Parking jest tuż obok lokalu, bez problemu zaparkujesz nawet w sobotę po południu. Z mapą Google trafisz w 5 minut z dowolnego punktu Nowego Targu.'),
        ('Jakie są godziny otwarcia?',
         'Codziennie od 10:00. Poniedziałek–piątek do 19:00, sobota i niedziela do 20:00. Otwarte przez cały rok — nie zamykamy się w wakacje ani między świętami. Ostatnie wejście do bawialni przyjmujemy 30 minut przed zamknięciem.'),
        ('Czy do bawialni trzeba rezerwować termin?',
         'Nie. Wstęp do bawialni jest bez rezerwacji — wchodzisz, kupujesz bilet przy ladzie (25 zł za godzinę, 33 zł za 2 godziny, 45 zł cały dzień, 215 zł karnet miesięczny bez limitu). Rezerwacji wymagają tylko urodziny, warsztaty i wyjścia grupowe.'),
        ('Jak zarezerwować urodziny dla dziecka?',
         'Najszybciej telefonem na +48 693 766 049 — od razu sprawdzamy wolne daty i ustalamy pakiet (MINI 45 zł/os. albo STANDARD 74 zł/os.). Można też wypełnić krótki formularz na podstronie urodzin albo napisać na Instagramie. Pakiet STANDARD rezerwuje się 3-6 tygodni wcześniej, MINI często łapiemy w tym samym tygodniu.'),
        ('Czy macie email do firmowych zapytań i faktur?',
         'Email do faktur, ofert dla przedszkoli, szkół i firm dostaniesz po pierwszym kontakcie telefonicznym albo na Messengerze. Tak najszybciej trafiamy w to, czego potrzebujesz, bez maili, które giną w spamie.'),
        ('Czy można odwiedzić Bawisz bez dziecka, na samą kawę?',
         'Tak. Kawiarnia działa osobno — wpadasz, zamawiasz przy ladzie, siadasz przy stoliku. Kawa rzemieślnicza, domowe ciasta (też bezglutenowe i bez cukru), bez biletu do bawialni. Dobre miejsce na spotkanie albo godzinę pracy z laptopem między spotkaniami.'),
    ]
    row(t, "EYEBROW", "[ FAQ ]")
    row(t, "NAGŁÓWEK H2", "Najczęstsze pytania.", bold=True, size=11)
    for i, (q, a) in enumerate(kontakt_faq, 1):
        row(t, f"PYTANIE {i}", q, bold=True)
        row(t, f"ODPOWIEDŹ {i}", a)

    sec(t, "7. FINAL CTA — końcowy banner")
    row(t, "NAGŁÓWEK H2", "Pytanie albo wątpliwość?", bold=True, size=11)
    row(t, "OPIS", "Najszybciej telefonem — odpowiadamy od ręki w godzinach otwarcia. Z urodzinami, warsztatami i grupami pisz albo dzwoń wcześniej, terminy znikają szybciej niż się wydaje.")
    row(t, "PRZYCISK 1", "Zadzwoń · 693 766 049")
    row(t, "PRZYCISK 2", "Zarezerwuj urodziny")

    sec(t, "8. META — SEO")
    row(t, "META TITLE", "Kontakt — Bawisz Nowy Targ | telefon, mapa, godziny")
    row(t, "META DESCRIPTION", "Kontakt z Bawiszem w Nowym Targu — telefon +48 693 766 049, ul. Krzywa 19B, otwarte codziennie od 10:00. Napisz na Instagramie lub odwiedź nas osobiście.")

    save(doc, os.path.join(OUTDIR, "kontakt.docx"))


# ════════════════════════════════════════════════════════════════════════
# WSPÓLNE ELEMENTY (Navbar, Footer, Modal urodzinowy, status)
# ════════════════════════════════════════════════════════════════════════
def gen_wspolne():
    reset_alt()
    doc, t = new_doc("Wspólne elementy (Navbar, Footer, Modal)")

    sec(t, "1. NAVBAR — górne menu (na każdej podstronie)")
    row(t, "LOGO — alt", "BAWISZ")
    row(t, "LOGO — nazwa", "BAWISZ")
    row(t, "LOGO — podtytuł", "Bawialnia · Kawiarnia")
    row(t, "MENU 1", "Home")
    row(t, "MENU 2", "O nas")
    row(t, "MENU 3", "Cennik")
    row(t, "MENU 4", "Urodziny")
    row(t, "MENU 5", "Warsztaty")
    row(t, "MENU 6", "Dla przedszkoli")
    row(t, "MENU 7", "Kawiarnia")
    row(t, "MENU 8", "Kontakt")
    row(t, "STATUS — otwarte", "Otwarte · do 19:00 (lub 20:00 w weekendy)")
    row(t, "STATUS — zamknięte", "Zamknięte")
    row(t, "PRZYCISK CTA MOBILE 1", "Zarezerwuj urodziny")
    row(t, "PRZYCISK CTA MOBILE 2", "+48 693 766 049")

    sec(t, "2. FOOTER — stopka (na każdej podstronie)")
    row(t, "BRAND — slogan (linia 1)", "po")
    row(t, "BRAND — slogan (linia 2, kursywa)", "bawisz")
    row(t, "BRAND — slogan (linia 3)", "się?")
    row(t, "PRZYCISK CTA", "Zarezerwuj urodziny")
    row(t, "TELEFON CTA", "+48 693 766 049")
    row(t, "KOLUMNA 1 — nagłówek", "Bawisz", bold=True)
    row(t, "KOLUMNA 1 — opis (linia 1)", "Kreatywna sala zabaw i kawiarnia.")
    row(t, "KOLUMNA 1 — opis (linia 2)", "ul. Krzywa 19B, 34-400 Nowy Targ.")
    row(t, "KOLUMNA 2 — nagłówek", "Co u nas", bold=True)
    row(t, "KOLUMNA 2 — link 1", "Sala zabaw")
    row(t, "KOLUMNA 2 — link 2", "Cennik")
    row(t, "KOLUMNA 2 — link 3", "Warsztaty")
    row(t, "KOLUMNA 2 — link 4", "Urodziny")
    row(t, "KOLUMNA 3 — nagłówek", "Odwiedź", bold=True)
    row(t, "KOLUMNA 3 — link 1", "Godziny otwarcia")
    row(t, "KOLUMNA 3 — link 2", "Kontakt")
    row(t, "KOLUMNA 3 — link 3", "Nawigacja")
    row(t, "KOLUMNA 3 — link 4", "Kawiarnia")
    row(t, "KOLUMNA 4 — nagłówek", "Kontakt", bold=True)
    row(t, "KOLUMNA 4 — telefon", "+48 693 766 049")
    row(t, "KOLUMNA 4 — godziny (linia 1)", "Pn–Pt 10:00–19:00")
    row(t, "KOLUMNA 4 — godziny (linia 2)", "Sb–Nd 10:00–20:00")
    row(t, "STOPKA — copyright", "© 2026 Bawisz · Sala zabaw i kawiarnia · Nowy Targ")
    row(t, "STOPKA — design", "Design i wykonanie AI Solutions Design")

    sec(t, "3. MODAL REZERWACJI URODZIN — formularz 3-stopniowy")
    note(t, "Modal otwiera się po kliknięciu „Zarezerwuj urodziny\" w dowolnym miejscu. 3 kroki + ekran sukcesu.")
    row(t, "EYEBROW", "[ Urodziny w Bawiszu ]")
    row(t, "TYTUŁ — szablon", "Krok {1/2/3} z 3")
    row(t, "KROK 1 — pole 1 (label)", "Imię solenizanta/-ki")
    row(t, "KROK 1 — pole 1 (placeholder)", "np. Antek")
    row(t, "KROK 1 — pole 2 (label)", "Wiek")
    row(t, "KROK 1 — pole 3 (label)", "Liczba dzieci")
    row(t, "KROK 1 — pole 4 (label)", "Preferowana data")
    row(t, "KROK 2 — pakiet 1 nazwa", "MINI", bold=True)
    row(t, "KROK 2 — pakiet 1 cena", "45 zł / dziecko")
    row(t, "KROK 2 — pakiet 1 opis", "2h zabawy, dekoracje, poczęstunek dla dzieci, obsługa")
    row(t, "KROK 2 — pakiet 2 nazwa", "STANDARD", bold=True)
    row(t, "KROK 2 — pakiet 2 cena", "74 zł / dziecko")
    row(t, "KROK 2 — pakiet 2 opis", "2,5h, sala na wyłączność (od 10 dzieci), prezent dla solenizanta, motyw lasu lub kwiatów")
    row(t, "KROK 2 — pakiet 3 nazwa", "Pakiet dla rodziców", bold=True)
    row(t, "KROK 2 — pakiet 3 cena", "+55 zł / osoba")
    row(t, "KROK 2 — pakiet 3 opis", "Dodatek do MINI lub STANDARD: kawa, herbata, ciasto, przekąski dla dorosłych")
    row(t, "KROK 3 — pole 1 (label)", "Imię i nazwisko")
    row(t, "KROK 3 — pole 1 (placeholder)", "Anna Kowalska")
    row(t, "KROK 3 — pole 2 (label)", "Telefon")
    row(t, "KROK 3 — pole 2 (placeholder)", "+48 …")
    row(t, "KROK 3 — pole 3 (label)", "Email")
    row(t, "KROK 3 — pole 3 (placeholder)", "email@example.pl")
    row(t, "KROK 3 — pole 4 (label)", "Uwagi (alergie, motyw, niespodzianki)")
    row(t, "KROK 3 — pole 4 (placeholder)", "np. motyw leśny, alergia na orzechy")
    row(t, "PRZYCISK — wstecz", "← Wstecz")
    row(t, "PRZYCISK — dalej", "Dalej")
    row(t, "PRZYCISK — wyślij", "Wyślij zapytanie")
    row(t, "SUKCES — nagłówek", "Dzięki! Damy znać.", bold=True)
    row(t, "SUKCES — opis (linia 1)", "Odezwiemy się w ciągu 24h, żeby potwierdzić termin i ustalić szczegóły.")
    row(t, "SUKCES — opis (linia 2)", "Tymczasem upewnij się, że telefon jest naładowany.")
    row(t, "SUKCES — przycisk", "Zamknij")

    save(doc, os.path.join(OUTDIR, "wspolne.docx"))


if __name__ == '__main__':
    print(f"Generuję dokumenty Word w: {OUTDIR}\n")
    gen_home()
    gen_urodziny()
    gen_kawiarnia()
    gen_warsztaty()
    gen_o_nas()
    gen_oferta_grupowa()
    gen_kontakt()
    gen_wspolne()
    print(f"\n✅ Gotowe — 8 plików .docx w {OUTDIR}")
